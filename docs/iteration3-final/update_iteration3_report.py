from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT.parent / "细化迭代2第9组 马启凡 叶炳良 周子栋 庄子杰.docx"
ASSET_DIR = ROOT / "docs" / "iteration3-final" / "generated"


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def body_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8 if level == 2 else 6)
    run = p.add_run(text)
    set_run_font(run, 18 if level == 2 else 14, True)
    return p


def chapter_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 16)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, 10.5)
    return p


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Normal Table"
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_picture(doc, filename, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ASSET_DIR / filename), width=Cm(15.5))
    caption(doc, title)


def truncate_existing_chapter(doc):
    body = doc.element.body
    children = list(body)
    start = None
    for index, child in enumerate(children):
        if child.tag == qn("w:p"):
            texts = child.xpath(".//w:t")
            combined = "".join(t.text or "" for t in texts)
            if "第四章 细化迭代 3" in combined or "第四章 细化迭代3" in combined:
                start = index
                break
    if start is not None:
        for child in children[start:]:
            body.remove(child)


def main():
    doc = Document(str(DOCX))
    truncate_existing_chapter(doc)

    doc.add_page_break()
    chapter_heading(doc, "第四章 细化迭代 3")

    body_para(doc, "根据总体迭代计划，第四章对应细化迭代 3，提交时间为第14周周五晚。本阶段的课程制品重点是 UML State Machine Diagrams and Modeling，并在此基础上 refine the above models，汇总形成完整文档。与第二章侧重领域模型、系统顺序图和操作契约，第三章侧重 GoF 设计模式不同，本章的核心任务不是重新定义 ATM 系统，而是在已有需求、模型和代码实现基础上，把最终演示链路、状态迁移、设备约束、流水凭条和部署视图收束到同一套可验证的设计说明中。")
    body_para(doc, "从项目实现角度看，第三次迭代在第二次迭代已经完成取款、存款、转账和修改密码等核心交易闭环的基础上，继续补齐交易流水查询、交易凭条查询、ATM 设备状态、吐钞能力检查以及最终演示入口。前端新增 TransactionHistoryView、ReceiptView 和 DeviceStatusView，并在主菜单、余额页和交易结果页串联这些页面；后端新增 ReceiptController、DeviceController、AtmDevice、AtmDeviceMapper 和 DeviceService，并扩展 TransactionService 与 TransactionMapper，使流水、凭条和设备现金变化能够与真实交易记录对应。")

    heading(doc, "4.1 迭代依据与目标")
    body_para(doc, "本轮迭代的依据包括总体迭代计划图片、三次迭代任务安排、UML 作业方案、第三次迭代记录、前后端 review 文档以及当前代码实现。总体迭代计划明确要求本章完成状态机建模和既有模型完善；三次迭代任务安排进一步把本轮拆解为交易流水、凭条、设备状态、异常处理、测试、UML 终稿和答辩材料整理；实验要求和作业方案则要求报告按 UP 过程逐步推进，使每一轮模型都能回到真实业务和系统实现。")
    add_table(doc, ["资料来源", "对第四章的约束", "本文落实方式"], [
        ["总体迭代计划.png", "第四章细化迭代 3：状态机建模、模型完善和完整文档汇总", "以状态机为主图，同时补充组件、部署、演示流程和模型演进关系"],
        ["ATM_四人分工与三次迭代任务.md", "完成流水、凭条、设备状态、测试、联调和 UML 终稿", "按前端、交易、设备和文档四条线说明实现成果"],
        ["ATM_UML作业方案.md", "迭代三覆盖打印凭条、交易流水、异常处理、系统部署与演示", "将凭条、流水、异常和部署视图纳入第四章主体"],
        ["第三次迭代记录.md", "记录新增接口、验证结果和最终演示流程", "作为功能完成度、接口表和测试证据的主要事实来源"],
        ["当前代码实现", "后端服务、前端页面、接口契约必须与报告一致", "所有表述锚定真实类、接口、页面和测试命令"],
    ])
    body_para(doc, "因此，本章的阶段目标可以概括为：第一，补充 ATM 会话与业务执行过程的状态机模型，说明系统在待机、认证、业务选择、设备检查、交易处理、凭条展示、流水查询和退卡之间如何转换；第二，完善交易流水、交易凭条和设备状态相关设计，使其与第二次迭代的交易闭环衔接；第三，汇总最终演示、部署和测试验证材料，形成可用于课程提交和答辩说明的完整文档闭环。")

    heading(doc, "4.2 第三次迭代功能完成情况")
    body_para(doc, "第三次迭代并没有推翻前两次迭代已经形成的架构，而是在原有前后端分离结构上进行增量完善。前端继续以 Vue 3、Vue Router、Pinia、Element Plus 和 Axios 为基础，通过 atm.js 统一封装接口调用；后端继续采用 Spring Boot、Controller、Service、Mapper 和统一 Result 响应结构，所有新增功能仍以 sessionId 为会话主字段，避免接口命名在最后阶段出现分叉。")
    add_table(doc, ["功能点", "前端实现", "后端实现", "设计意义"], [
        ["交易流水", "TransactionHistoryView 支持分页展示交易时间、编号、类型、金额和状态", "TransactionService.getTransactionHistory 按当前账户分页查询", "把交易结果沉淀为可追溯记录，支撑审计和凭条入口"],
        ["交易凭条", "ReceiptView 支持从交易结果、流水记录或手动输入编号进入", "ReceiptController 调用 getReceipt，并按 sessionId 限定账户范围", "体现凭条不是孤立页面，而是交易记录的受控视图"],
        ["设备状态", "DeviceStatusView 展示 ATM 编号、位置、运行状态和可用现金", "DeviceService.getStatus 读取 atm_device 初始化设备 ATM001", "把 ATM 设备从背景资源提升为可建模对象"],
        ["吐钞检查", "取款前调用 checkCashAvailability，现金不足时提前提示", "DeviceService.ensureCashAvailable 和 dispenseCash 参与取款事务流程", "使账户余额和设备现金能力共同决定取款是否可执行"],
        ["最终演示入口", "主菜单、余额页和交易成功页串联流水、凭条、设备页面", "接口与 README、OpenAPI、HTTP 示例保持一致", "形成从登录到退卡的完整演示路径"],
    ])
    body_para(doc, "从功能边界看，本轮的重点在于完成“交易之后发生什么”和“ATM 设备是否具备执行能力”这两个前两轮尚未完全落地的问题。交易流水和凭条回答了交易结果如何保存、查询和展示；设备状态和吐钞检查回答了取款并非只依赖账户余额，还必须受到终端运行状态和现金余量约束。")

    heading(doc, "4.3 会话与业务状态机建模")
    body_para(doc, "总体迭代计划要求本章重点完成状态机建模。ATM 系统天然具有明显的状态特征：设备从待机进入认证，认证通过后进入业务选择，客户选择不同业务后进入交易处理或设备检查，交易成功后可以查看凭条和流水，修改密码会导致当前会话失效，退卡后系统重新回到待机状态。将这些状态显式建模，可以比单纯的流程图更清楚地说明“系统当前允许什么操作、异常后回到哪里、会话何时失效”。")
    add_picture(doc, "atm-session-state.png", "图 4-1 细化迭代 3 ATM 会话与业务状态机")
    body_para(doc, "图 4-1 中，sessionId 是连接认证状态和后续业务状态的关键条件。只有认证成功并持有有效 sessionId 后，客户才能进入余额查询、交易办理、流水查询和凭条查看等状态。修改密码成功后，当前 sessionId 会失效，系统必须回到重新登录路径；交易异常和设备异常则不直接结束会话，而是进入异常提示并允许客户返回业务选择重新操作。这一处理方式既符合 ATM 用户体验，也与后端 SessionValidator 和统一异常处理的实现相对应。")

    heading(doc, "4.4 流水与凭条的受控查询模型")
    body_para(doc, "交易流水和交易凭条是第三次迭代最能体现最终演示完整性的功能。第二次迭代已经能够生成取款、存款和转账交易记录，但如果没有流水页面和凭条页面，客户只能在交易完成的瞬间看到结果，无法从历史视角追溯交易。本轮通过交易流水分页接口和凭条查询接口，把交易记录转化为可查询、可展示、可验证的业务对象。")
    add_picture(doc, "receipt-history-sequence.png", "图 4-2 交易流水与凭条查询顺序图")
    body_para(doc, "图 4-2 强调了一个重要边界：凭条查询并不只依赖 transactionId。后端在 getReceipt 中会先通过 sessionId 定位当前账户，再查询 transactionId 对应的交易记录；如果交易不存在，或者交易存在但不属于当前会话账户，系统都返回凭条不存在。这样处理可以避免客户通过猜测交易编号读取其他账户的交易信息，也使凭条功能符合第一章和第二章中关于账户边界、会话控制和异常处理的设计原则。")

    heading(doc, "4.5 设备状态与取款现金联动")
    body_para(doc, "在初始用例中，ATM 吐钞模块和打印模块已经被识别为会影响业务结果的边界资源。第三次迭代把设备状态进一步落实到代码层：atm_device 表保存 ATM001 的设备编号、位置、运行状态和可用现金；DeviceService 提供设备状态查询、吐钞能力检查、取款前现金确认和取款成功后的现金扣减。由此，取款流程从“账户余额是否足够”扩展为“账户余额和设备现金能力是否同时满足”。")
    add_picture(doc, "device-cash-flow.png", "图 4-3 取款设备现金联动流程")
    body_para(doc, "图 4-3 反映了本轮的一个关键扩展点：取款交易中，系统先完成金额规则、单次限额、单日限额和账户余额检查，再通过设备服务确认 ATM 是否运行且现金充足。只有这些条件全部满足后，系统才创建 PENDING 状态的交易流水，随后扣减账户余额、扣减设备现金，并最终将交易标记为 SUCCESS。存款和转账不改变设备现金，这一边界避免了把“账户资金变化”和“ATM 机内现金变化”混为一谈。")

    heading(doc, "4.6 组件与部署视图完善")
    body_para(doc, "第三次迭代还承担 refine the above models 的任务，因此本章需要把前两章已出现的领域对象、控制对象、服务对象和数据对象进一步组织成最终组件视图。与第二次迭代仅围绕交易服务内部模式不同，本轮组件视图同时覆盖显示模块、接口封装、认证账户模块、交易模块、凭条模块、设备模块、统一异常处理和数据层表结构。")
    add_picture(doc, "final-component-deployment.png", "图 4-4 第三次迭代组件与部署视图")
    body_para(doc, "图 4-4 说明了最终系统的部署和职责划分。用户侧通过浏览器或 ATM 模拟终端访问 Vue 前端页面，前端通过 atm.js 调用后端 REST 接口；后端由认证账户、交易、凭条和设备模块组成，并通过 Mapper 访问 customer、bank_card、account、transaction_record 和 atm_device 等数据表；统一异常处理负责把业务异常转换为前端可识别的 Result 响应。该视图既能支撑最终部署说明，也能帮助答辩时解释前后端分离架构和模块边界。")

    heading(doc, "4.7 最终演示流程与验收证据")
    body_para(doc, "第三次迭代的最终演示不应只展示单个页面，而应覆盖从登录、设备检查、交易办理、凭条查看、流水查询到退卡退出的完整路径。当前项目已经在 README 和第三次迭代记录中整理了推荐演示顺序：启动前后端真实联调模式，使用演示卡号登录，进入设备状态页执行吐钞能力检查，办理一笔取款、存款或转账，随后查看凭条和流水，最后退卡退出。")
    add_picture(doc, "final-demo-flow.png", "图 4-5 第三次迭代最终演示流程")
    add_table(doc, ["验证项", "验证方式", "当前结果"], [
        ["后端功能", "在 atm-server-auth 目录执行 .\\mvnw.cmd test", "Tests run: 19, Failures: 0, Errors: 0, Skipped: 0，BUILD SUCCESS"],
        ["前端构建", "在 frontend 目录执行 npm run build", "生产构建通过，存在 Element Plus 体积相关 warning，不阻塞演示"],
        ["流水与凭条", "登录后办理交易，再查询 /transactions/history 和 /receipts/{transactionId}", "流水按当前账户分页返回，凭条需匹配当前 sessionId"],
        ["设备能力", "调用 /api/atm/device/status 和 /api/atm/device/cash-check", "能返回 ATM001 状态和现金可用性，取款成功后扣减设备现金"],
        ["会话边界", "修改密码后继续使用旧 sessionId 访问交易类接口", "当前会话失效，需要重新登录"],
    ])
    body_para(doc, "上述验证结果说明，第三次迭代已经从单一功能补齐推进到整体演示可复现。尤其是流水、凭条和设备状态三个页面接入主菜单后，客户可以在同一套界面中完成“办理交易 - 查看凭条 - 追溯流水 - 检查设备”的闭环操作，符合课程最终演示对完整性的要求。")

    heading(doc, "4.8 模型完善与文档汇总")
    body_para(doc, "从 UP 过程角度看，第四章不是孤立的新章节，而是对前三章模型的收束。第一章建立系统边界和用例基础，第二章通过领域模型、系统顺序图和操作契约细化核心流程，第三章把核心交易闭环与 GoF 模式结合起来，本章则在这些模型之上补充状态机、组件部署视图、设备模型和最终演示证据。")
    add_picture(doc, "model-refinement-map.png", "图 4-6 从前三章到细化迭代 3 的模型完善关系")
    add_table(doc, ["前三章已有基础", "第四章完善内容", "完善后的价值"], [
        ["用例模型与关键规约", "补充状态机，说明用例执行时系统状态如何迁移", "让登录、交易、凭条、流水和退卡之间的关系更清晰"],
        ["领域模型和操作契约", "补充交易记录、凭条访问和设备现金约束", "把状态变化和数据变化对应到真实对象"],
        ["交易模块 GoF 设计", "扩展到凭条、流水和设备模块的协作", "说明第二次迭代设计如何支撑第三次迭代扩展"],
        ["前后端接口契约", "统一整理新增接口、演示流程和测试结果", "保证代码、OpenAPI、README 和报告文档一致"],
    ])
    body_para(doc, "本轮文档汇总的重点不是堆叠更多图，而是让每一张图都有明确作用：状态机解释系统状态约束，顺序图解释流水和凭条的访问控制，流程图解释设备现金如何参与取款事务，组件部署图解释系统运行结构，演示流程图解释最终答辩时如何复现第三次迭代成果。")

    heading(doc, "4.9 亮点、拓展点与阶段结论")
    body_para(doc, "第三次迭代的第一个亮点是公开契约继续保持统一。前端、OpenAPI、README、HTTP 示例和后端实现均以 sessionId、transactionId、targetAccountNo 和 /api/atm/accounts/* 为主口径，token 仅作为兼容字段存在。这使得报告中描述的接口、前端调用和后端代码能够互相印证，避免最终阶段出现字段混乱。")
    body_para(doc, "第二个亮点是交易凭条和交易流水都具备会话账户隔离。系统不是简单按 transactionId 返回凭条，而是必须结合 sessionId 验证当前账户；流水查询也只返回当前账户记录。这一设计虽然属于课程项目中的轻量级访问控制，但已经体现了 ATM 业务中“客户只能查看本人交易信息”的基本安全要求。")
    body_para(doc, "第三个亮点是设备状态被纳入真实业务流程。取款前检查 ATM 现金是否充足，取款成功后扣减设备可用现金，使设备不再只是图中的背景节点，而成为影响交易结果的业务资源。后续如果继续拓展，可以在此基础上加入设备维护、现金补充、打印机状态、设备故障告警和管理员后台等功能。")
    body_para(doc, "综合来看，细化迭代 3 已经完成总体迭代计划中对状态机建模、模型完善和完整文档汇总的要求，也完成三次迭代任务中流水、凭条、设备状态、异常处理、测试和演示流程整理等主要目标。至此，ATM 系统从初始阶段的需求与用例识别，逐步演进到细化迭代 1 的模型基线、细化迭代 2 的核心交易闭环，再到细化迭代 3 的最终演示和文档收束，形成了较完整的软件分析设计与建模交付成果。")

    doc.save(str(DOCX))


if __name__ == "__main__":
    main()
